"""
NETRIX professional report generation (PDF, Excel, Word, CSV).
PDF uses Tahoma-compatible sans-serif (Liberation Sans / system Tahoma),
black 12pt body text, and detailed topology + UML graphics.
"""
import io
import os
from datetime import datetime
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Preformatted, PageBreak, Flowable, HRFlowable, KeepTogether,
)
from reportlab.graphics.shapes import (
    Drawing, Rect, String, Line, Circle, Polygon, Group, Ellipse,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── Font registration (Tahoma preferred, Liberation Sans fallback) ─────────

def _register_fonts():
    """Register Tahoma if present, else Liberation Sans / DejaVu as 'Tahoma'."""
    candidates = [
        # Windows
        (r'C:\Windows\Fonts\tahoma.ttf', r'C:\Windows\Fonts\tahomabd.ttf'),
        # Wine / cross-platform
        (os.path.expanduser('~/.wine/drive_c/windows/fonts/tahoma.ttf'),
         os.path.expanduser('~/.wine/drive_c/windows/fonts/tahomabd.ttf')),
        # Liberation Sans (metric-compatible business sans)
        ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
         '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'),
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
    ]
    for regular, bold in candidates:
        if os.path.isfile(regular) and os.path.isfile(bold):
            try:
                pdfmetrics.registerFont(TTFont('Tahoma', regular))
                pdfmetrics.registerFont(TTFont('Tahoma-Bold', bold))
                return 'Tahoma', 'Tahoma-Bold'
            except Exception:
                continue
    return 'Helvetica', 'Helvetica-Bold'


FONT_REG, FONT_BOLD = _register_fonts()

# Professional palette
BLACK = colors.HexColor('#000000')
DARK = colors.HexColor('#1a1a1a')
NAVY = colors.HexColor('#0f2744')
BLUE = colors.HexColor('#1e4d8c')
BLUE_LIGHT = colors.HexColor('#e8eef6')
GREY_LINE = colors.HexColor('#c8c8c8')
GREY_ROW = colors.HexColor('#f5f5f5')
WHITE = colors.white
GREEN = colors.HexColor('#0d6b4c')
ORANGE = colors.HexColor('#c45c00')
PURPLE = colors.HexColor('#5b2d8e')


def _styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='RTitle',
        fontName=FONT_BOLD,
        fontSize=18,
        textColor=BLACK,
        alignment=TA_CENTER,
        spaceAfter=6,
        leading=22,
    ))
    styles.add(ParagraphStyle(
        name='RSubtitle',
        fontName=FONT_REG,
        fontSize=12,
        textColor=BLACK,
        alignment=TA_CENTER,
        spaceAfter=4,
        leading=16,
    ))
    styles.add(ParagraphStyle(
        name='RH1',
        fontName=FONT_BOLD,
        fontSize=14,
        textColor=BLACK,
        spaceBefore=16,
        spaceAfter=8,
        leading=18,
        borderPadding=3,
    ))
    styles.add(ParagraphStyle(
        name='RH2',
        fontName=FONT_BOLD,
        fontSize=12,
        textColor=BLACK,
        spaceBefore=12,
        spaceAfter=6,
        leading=16,
    ))
    styles.add(ParagraphStyle(
        name='RBody',
        fontName=FONT_REG,
        fontSize=12,
        textColor=BLACK,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
        leading=16,
    ))
    styles.add(ParagraphStyle(
        name='RCaption',
        fontName=FONT_REG,
        fontSize=10,
        textColor=DARK,
        alignment=TA_CENTER,
        spaceBefore=4,
        spaceAfter=10,
        leading=13,
    ))
    styles.add(ParagraphStyle(
        name='RSmall',
        fontName=FONT_REG,
        fontSize=10,
        textColor=BLACK,
        leading=13,
    ))
    styles.add(ParagraphStyle(
        name='RCode',
        fontName='Courier',
        fontSize=8,
        textColor=BLACK,
        backColor=colors.HexColor('#f4f4f4'),
        leading=10,
        leftIndent=4,
        rightIndent=4,
    ))
    styles.add(ParagraphStyle(
        name='RFooter',
        fontName=FONT_REG,
        fontSize=9,
        textColor=DARK,
        alignment=TA_CENTER,
    ))
    return styles


def _header_footer(canvas, doc):
    canvas.saveState()
    page_w, page_h = A4
    # Top bar
    canvas.setFillColor(NAVY)
    canvas.rect(0, page_h - 28, page_w, 28, fill=1, stroke=0)
    canvas.setFillColor(WHITE)
    canvas.setFont(FONT_BOLD, 10)
    canvas.drawString(40, page_h - 18, 'NETRIX  |  Enterprise Network Planning Report')
    canvas.setFont(FONT_REG, 9)
    canvas.drawRightString(page_w - 40, page_h - 18, datetime.utcnow().strftime('%Y-%m-%d'))
    # Bottom bar
    canvas.setFillColor(NAVY)
    canvas.rect(0, 0, page_w, 24, fill=1, stroke=0)
    canvas.setFillColor(WHITE)
    canvas.setFont(FONT_REG, 9)
    canvas.drawString(40, 9, 'Confidential – For authorised use only')
    canvas.drawRightString(page_w - 40, 9, f'Page {doc.page}')
    # Thin accent line under header
    canvas.setStrokeColor(BLUE)
    canvas.setLineWidth(2)
    canvas.line(40, page_h - 32, page_w - 40, page_h - 32)
    canvas.restoreState()


# ─── Diagrams ───────────────────────────────────────────────────────────────

class DiagramFlowable(Flowable):
    def __init__(self, drawing: Drawing, caption: str = ''):
        super().__init__()
        self.drawing = drawing
        self.caption = caption
        self.width = drawing.width
        self.height = drawing.height + (16 if caption else 0)

    def draw(self):
        self.drawing.drawOn(self.canv, 0, 16 if self.caption else 0)
        if self.caption:
            self.canv.setFont(FONT_REG, 9)
            self.canv.setFillColor(DARK)
            self.canv.drawCentredString(self.width / 2, 2, self.caption)


def _draw_topology_diagram(project: dict, generated: dict) -> Drawing:
    vlsm = generated.get('vlsm') or []
    depts = [v.get('dept', f'D{i}') for i, v in enumerate(vlsm[:8])]
    routers = project.get('routers') or [{'name': project.get('router_name') or 'R1', 'role': 'edge'}]
    switches = project.get('switches') or [{'name': project.get('switch_name') or 'S1', 'role': 'access'}]
    internet = project.get('internet') or {'enabled': True, 'name': 'Internet'}
    if not depts:
        depts = ['LAN']

    width, height = 500, 320
    d = Drawing(width, height)

    # Card background
    d.add(Rect(0, 0, width, height, fillColor=colors.HexColor('#fafbfc'),
               strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Rect(0, height - 26, width, 26, fillColor=NAVY, strokeColor=None))
    d.add(String(12, height - 17, 'Figure – Network Topology',
                 fontSize=11, fillColor=WHITE, fontName=FONT_BOLD))

    y_net, y_r, y_s, y_d = height - 70, height - 130, height - 195, 48

    # Internet cloud
    if internet.get('enabled', True):
        cx = width / 2
        for ox, oy, r in [(-20, 0, 14), (0, 5, 16), (20, 0, 14), (-8, -6, 11), (8, -6, 11)]:
            d.add(Circle(cx + ox, y_net + oy, r,
                         fillColor=colors.HexColor('#dbeafe'),
                         strokeColor=BLUE, strokeWidth=1.2))
        d.add(String(cx - 22, y_net - 3, internet.get('name') or 'Internet',
                     fontSize=9, fillColor=BLACK, fontName=FONT_BOLD))
        d.add(String(cx - 16, y_net - 15, 'WAN / ISP', fontSize=7, fillColor=DARK))

    # Routers
    n_r = len(routers)
    r_sp = min(95, (width - 50) / max(n_r, 1))
    r_start = (width - r_sp * (n_r - 1)) / 2 if n_r > 1 else width / 2
    router_xs = []
    for i, r in enumerate(routers):
        x = r_start + i * r_sp if n_r > 1 else width / 2
        router_xs.append(x)
        if internet.get('enabled', True):
            d.add(Line(width / 2, y_net - 16, x, y_r + 14,
                       strokeColor=BLUE, strokeWidth=1.3))
        d.add(Rect(x - 40, y_r - 14, 80, 28, fillColor=BLUE,
                   strokeColor=NAVY, strokeWidth=1, rx=4, ry=4))
        d.add(String(x - 26, y_r - 2, (r.get('name') or f'R{i+1}')[:12],
                     fontSize=9, fillColor=WHITE, fontName=FONT_BOLD))
        d.add(String(x - 18, y_r - 26, (r.get('role') or 'router')[:10],
                     fontSize=7, fillColor=DARK))

    # Switches
    n_s = len(switches)
    s_sp = min(95, (width - 50) / max(n_s, 1))
    s_start = (width - s_sp * (n_s - 1)) / 2 if n_s > 1 else width / 2
    switch_xs = []
    for i, sw in enumerate(switches):
        x = s_start + i * s_sp if n_s > 1 else width / 2
        switch_xs.append(x)
        rx = router_xs[min(i, len(router_xs) - 1)]
        d.add(Line(rx, y_r - 14, x, y_s + 14, strokeColor=GREEN, strokeWidth=1.3))
        d.add(Rect(x - 40, y_s - 14, 80, 28, fillColor=GREEN,
                   strokeColor=colors.HexColor('#064e3b'), strokeWidth=1, rx=4, ry=4))
        d.add(String(x - 26, y_s - 2, (sw.get('name') or f'S{i+1}')[:12],
                     fontSize=9, fillColor=WHITE, fontName=FONT_BOLD))
        d.add(String(x - 20, y_s - 26, (sw.get('role') or 'access')[:10],
                     fontSize=7, fillColor=DARK))

    # Departments
    n_d = len(depts)
    d_sp = min(58, (width - 36) / max(n_d, 1))
    d_start = (width - d_sp * (n_d - 1)) / 2 if n_d > 1 else width / 2
    for i, name in enumerate(depts):
        x = d_start + i * d_sp if n_d > 1 else width / 2
        sx = switch_xs[min(i % max(len(switch_xs), 1), len(switch_xs) - 1)]
        d.add(Line(sx, y_s - 14, x, y_d + 16, strokeColor=GREY_LINE, strokeWidth=1))
        d.add(Rect(x - 26, y_d - 11, 52, 26, fillColor=colors.HexColor('#eef2f7'),
                   strokeColor=DARK, strokeWidth=0.8, rx=3, ry=3))
        label = (name[:8] + '…') if len(name) > 8 else name
        d.add(String(x - 22, y_d - 1, label, fontSize=7, fillColor=BLACK, fontName=FONT_REG))

    proto = project.get('routing_protocol') or 'OSPF'
    base = project.get('base_network') or ''
    d.add(String(10, 8,
                 f'{proto}  ·  {base}  ·  {n_r} router(s)  ·  {n_s} switch(es)  ·  {n_d} segment(s)',
                 fontSize=8, fillColor=DARK, fontName=FONT_REG))
    return d


def _draw_uml_usecase(project: dict) -> Drawing:
    width, height = 500, 250
    d = Drawing(width, height)
    d.add(Rect(0, 0, width, height, fillColor=WHITE, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Rect(0, height - 24, width, 24, fillColor=NAVY, strokeColor=None))
    d.add(String(12, height - 16, 'Figure – UML Use Case Diagram',
                 fontSize=11, fillColor=WHITE, fontName=FONT_BOLD))

    # Actor Engineer
    ax, ay = 48, height / 2 - 5
    d.add(Circle(ax, ay + 28, 9, fillColor=WHITE, strokeColor=BLACK, strokeWidth=1.5))
    d.add(Line(ax, ay + 19, ax, ay - 2, strokeColor=BLACK, strokeWidth=1.5))
    d.add(Line(ax - 11, ay + 10, ax + 11, ay + 10, strokeColor=BLACK, strokeWidth=1.5))
    d.add(Line(ax, ay - 2, ax - 9, ay - 20, strokeColor=BLACK, strokeWidth=1.5))
    d.add(Line(ax, ay - 2, ax + 9, ay - 20, strokeColor=BLACK, strokeWidth=1.5))
    d.add(String(ax - 18, ay - 34, 'Engineer', fontSize=9, fillColor=BLACK, fontName=FONT_BOLD))

    # System boundary
    d.add(Rect(105, 28, 290, 185, fillColor=colors.HexColor('#f8fafc'),
               strokeColor=BLUE, strokeWidth=1.5))
    d.add(String(200, 198, 'NETRIX System', fontSize=9, fillColor=BLUE, fontName=FONT_BOLD))

    cases = [
        'Create Project', 'Configure VLSM', 'Allocate VLANs',
        'Generate Configs', 'Export Reports', 'Validate Design',
    ]
    for i, name in enumerate(cases):
        y = 175 - i * 24
        d.add(Line(ax + 10, ay, 120, y + 7, strokeColor=GREY_LINE, strokeWidth=0.8))
        d.add(Ellipse(180, y + 8, 55, 11, fillColor=BLUE_LIGHT, strokeColor=BLUE, strokeWidth=1))
        d.add(String(140, y + 5, name, fontSize=8, fillColor=BLACK, fontName=FONT_REG))

    # Admin actor
    bx = 450
    d.add(Circle(bx, ay + 28, 9, fillColor=WHITE, strokeColor=PURPLE, strokeWidth=1.5))
    d.add(Line(bx, ay + 19, bx, ay - 2, strokeColor=PURPLE, strokeWidth=1.5))
    d.add(Line(bx - 11, ay + 10, bx + 11, ay + 10, strokeColor=PURPLE, strokeWidth=1.5))
    d.add(Line(bx, ay - 2, bx - 9, ay - 20, strokeColor=PURPLE, strokeWidth=1.5))
    d.add(Line(bx, ay - 2, bx + 9, ay - 20, strokeColor=PURPLE, strokeWidth=1.5))
    d.add(String(bx - 14, ay - 34, 'Admin', fontSize=9, fillColor=PURPLE, fontName=FONT_BOLD))
    d.add(Line(bx - 10, ay, 250, 48, strokeColor=colors.HexColor('#c4b5fd'), strokeWidth=0.8))
    return d


def _draw_uml_activity(project: dict) -> Drawing:
    width, height = 500, 200
    d = Drawing(width, height)
    d.add(Rect(0, 0, width, height, fillColor=WHITE, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Rect(0, height - 24, width, 24, fillColor=NAVY, strokeColor=None))
    d.add(String(12, height - 16, 'Figure – UML Activity Flow',
                 fontSize=11, fillColor=WHITE, fontName=FONT_BOLD))

    steps = [
        ('1. Details', BLUE),
        ('2. Depts', PURPLE),
        ('3. VLSM', colors.HexColor('#0369a1')),
        ('4. VLAN', GREEN),
        ('5. Configs', ORANGE),
        ('6. Reports', colors.HexColor('#b91c1c')),
    ]
    box_w, box_h = 68, 34
    gap = 10
    total = len(steps) * box_w + (len(steps) - 1) * gap
    start = (width - total) / 2
    y = height / 2 - box_h / 2 - 5
    for i, (label, color) in enumerate(steps):
        x = start + i * (box_w + gap)
        d.add(Rect(x, y, box_w, box_h, fillColor=color, strokeColor=None, rx=4, ry=4))
        d.add(String(x + 8, y + 12, label, fontSize=8, fillColor=WHITE, fontName=FONT_BOLD))
        if i < len(steps) - 1:
            ax = x + box_w + gap
            d.add(Line(x + box_w, y + box_h / 2, ax - 4, y + box_h / 2,
                       strokeColor=DARK, strokeWidth=1.4))
            d.add(Polygon(
                [ax, y + box_h / 2, ax - 5, y + box_h / 2 + 4, ax - 5, y + box_h / 2 - 4],
                fillColor=DARK, strokeColor=None,
            ))
    d.add(String(12, 10,
                 f'Routing: {project.get("routing_protocol", "OSPF")}  ·  Base: {project.get("base_network", "")}',
                 fontSize=8, fillColor=DARK, fontName=FONT_REG))
    return d


def _draw_uml_class(project: dict, generated: dict) -> Drawing:
    width, height = 500, 270
    d = Drawing(width, height)
    d.add(Rect(0, 0, width, height, fillColor=WHITE, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Rect(0, height - 24, width, 24, fillColor=NAVY, strokeColor=None))
    d.add(String(12, height - 16, 'Figure – UML Component / Class View',
                 fontSize=11, fillColor=WHITE, fontName=FONT_BOLD))

    def box(x, y, w, h, title, lines, fill, stroke):
        d.add(Rect(x, y, w, h, fillColor=colors.HexColor(fill),
                   strokeColor=colors.HexColor(stroke), strokeWidth=1.2, rx=3, ry=3))
        d.add(Rect(x, y + h - 18, w, 18, fillColor=colors.HexColor(stroke), strokeColor=None))
        d.add(String(x + 6, y + h - 13, title, fontSize=8, fillColor=WHITE, fontName=FONT_BOLD))
        for i, line in enumerate(lines[:4]):
            d.add(String(x + 6, y + h - 32 - i * 12, str(line)[:26],
                         fontSize=7, fillColor=BLACK, fontName=FONT_REG))

    routers = project.get('routers') or [{'name': project.get('router_name') or 'R1'}]
    switches = project.get('switches') or [{'name': project.get('switch_name') or 'S1'}]
    vlsm = generated.get('vlsm') or []

    box(16, 150, 145, 85, '«Router» Edge',
        [f'+ {r.get("name")}' for r in routers] + [f'proto: {project.get("routing_protocol", "OSPF")}'],
        '#eff6ff', '#1e4d8c')
    box(178, 150, 145, 85, '«Switch» Fabric',
        [f'+ {s.get("name")}' for s in switches] + ['mode: access/trunk'],
        '#ecfdf5', '#0d6b4c')
    box(340, 150, 145, 85, '«Internet» Link',
        ['+ WAN /30 peer', f'on: {project.get("internet", {}).get("enabled", True)}',
         f'base: {project.get("base_network", "")}'],
        '#f5f3ff', '#5b2d8e')
    box(90, 35, 155, 90, '«Subnet» VLSM',
        [f'{v.get("dept")}: {v.get("network")}{v.get("prefix")}' for v in vlsm[:4]] or ['(none)'],
        '#fff7ed', '#c45c00')
    box(280, 35, 155, 90, '«VLAN» Allocation',
        [f'VLAN {v.get("vlan_id")}: {v.get("dept")}' for v in vlsm[:4]] or ['(none)'],
        '#f0f9ff', '#0369a1')

    d.add(Line(161, 175, 178, 175, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Line(323, 175, 340, 175, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Line(80, 150, 140, 125, strokeColor=GREY_LINE, strokeWidth=1))
    d.add(Line(250, 150, 340, 125, strokeColor=GREY_LINE, strokeWidth=1))
    return d


def _table_style(header=True):
    cmds = [
        ('FONTNAME', (0, 0), (-1, -1), FONT_REG),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TEXTCOLOR', (0, 0), (-1, -1), BLACK),
        ('GRID', (0, 0), (-1, -1), 0.6, GREY_LINE),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [WHITE, GREY_ROW]),
    ]
    if header:
        cmds += [
            ('BACKGROUND', (0, 0), (-1, 0), NAVY),
            ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
            ('FONTNAME', (0, 0), (-1, 0), FONT_BOLD),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
        ]
    return TableStyle(cmds)


# ─── PDF Report ─────────────────────────────────────────────────────────────


def _build_executive_summary(project: dict, generated: dict) -> list:
    """Return list of body paragraph HTML strings for the executive summary."""
    pname = project.get('project_name') or 'Network Project'
    company = project.get('company_name') or 'Organisation'
    base = project.get('base_network') or 'N/A'
    routing = project.get('routing_protocol') or 'OSPF'
    vlsm = generated.get('vlsm') or []
    vlan = generated.get('vlan') or []
    ipv4 = generated.get('ipv4') or []
    ipv6 = generated.get('ipv6') or []
    design = generated.get('design') or {}
    ip_version = (design.get('ip_version') or 'ipv4').upper()
    net_class = design.get('network_class') or 'C'
    topo = design.get('topology_type') or 'hierarchical'

    total_usable = 0
    for v in vlsm:
        try:
            total_usable += int(v.get('usable') or 0)
        except (TypeError, ValueError):
            pass
    depts = [v.get('dept') for v in vlsm if v.get('dept')]
    if not depts:
        depts = [v.get('dept') for v in ipv6 if v.get('dept') and v.get('vlan_id')]

    paras = [
        (
            f'This executive summary presents the key outcomes of the <b>{pname}</b> network design '
            f'prepared for <b>{company}</b>. The plan covers addressing ({ip_version}), VLAN segmentation, '
            f'topology style <b>{topo}</b>, private network class <b>{net_class}</b>, and Cisco IOS configurations '
            f'produced by the NETRIX Enterprise Network Planning Framework.'
        ),
        (
            f'<b>Scope.</b> The design uses base network <b>{base}</b> with routing protocol '
            f'<b>{routing}</b>. <b>{len(depts)}</b> departmental segment(s) were allocated. '
            f'IPv4 plan rows: <b>{len(ipv4)}</b>; IPv6 ULA plan rows: <b>{len(ipv6)}</b>; '
            f'VLANs defined: <b>{len(vlan)}</b>.'
        ),
    ]
    if total_usable:
        paras.append(
            f'<b>IPv4 capacity.</b> Approximate usable IPv4 host capacity across VLSM blocks is '
            f'<b>{total_usable}</b> addresses (excluding network/broadcast per subnet).'
        )
    if ipv6:
        samples = ', '.join(
            f'{r.get("dept")} ({r.get("network")})' for r in ipv6[:3] if r.get('network')
        )
        paras.append(
            f'<b>IPv6 addressing.</b> Unique Local Address (ULA) prefixes under '
            f'<font face="Courier">fd00:9e71::/32</font> were assigned on a /64 boundary per segment. '
            f'Examples: {samples}.'
        )
    if vlsm:
        largest = max(vlsm, key=lambda x: int(x.get('usable') or 0) if str(x.get('usable', '')).isdigit() else 0, default=None)
        if largest:
            paras.append(
                f'<b>Addressing highlights.</b> Subnets were assigned largest-first where applicable. '
                f'A primary segment is <b>{largest.get("dept")}</b> on <b>{largest.get("network")}</b> '
                f'(gateway {largest.get("gateway", "N/A")}).'
            )
    paras.append(
        '<b>Recommendation.</b> Validate connectivity and routing in Cisco Packet Tracer using the '
        'exported lab guide and device configurations, then implement under formal change control.'
    )
    return paras



def generate_pdf_report(project: dict, generated: dict) -> bytes:
    # Ensure optional keys exist
    generated = dict(generated or {})
    generated.setdefault("ipv6", [])
    generated.setdefault("design", {})
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.65 * inch, bottomMargin=0.55 * inch,
    )
    S = _styles()
    story = []

    # Cover / title block
    story.append(Spacer(1, 8))
    story.append(Paragraph('NETRIX', S['RTitle']))
    story.append(Paragraph('Enterprise Network Planning Report', S['RSubtitle']))
    story.append(HRFlowable(width='100%', thickness=2, color=NAVY, spaceAfter=8, spaceBefore=4))
    story.append(Paragraph(
        f"<b>Project:</b> {project.get('project_name', '')} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>Company:</b> {project.get('company_name', '')}",
        S['RSubtitle'],
    ))
    story.append(Paragraph(
        f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        S['RCaption'],
    ))

    # 1. Executive Summary
    story.append(Paragraph('1. Executive Summary', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    for para in _build_executive_summary(project, generated):
        story.append(Paragraph(para, S['RBody']))
        story.append(Spacer(1, 4))

    # Key metrics – expanded
    vlsm = generated.get('vlsm') or []
    vlan = generated.get('vlan') or []
    ipv4 = generated.get('ipv4') or []
    validation = generated.get('validation') or []
    routers = project.get('routers') or [{'name': project.get('router_name') or 'R1'}]
    switches = project.get('switches') or [{'name': project.get('switch_name') or 'S1'}]
    internet_on = (project.get('internet') or {}).get('enabled', True)
    total_usable = sum(int(v.get('usable') or 0) for v in vlsm)
    total_checks = len(validation)
    checks_pass = sum(1 for v in validation if v.get('pass'))
    checks_fail = total_checks - checks_pass
    largest = max(vlsm, key=lambda v: int(v.get('usable') or 0)) if vlsm else {}
    smallest = min(vlsm, key=lambda v: int(v.get('usable') or 0)) if vlsm else {}
    vlan_ids = sorted({int(v.get('vlan_id') or v.get('id') or 0) for v in (vlsm or vlan)})
    vlan_ids = [x for x in vlan_ids if x]

    story.append(Paragraph('<b>Key Metrics</b>', S['RH2']))
    # Row 1 – capacity & topology
    metrics1 = [
        ['Metric', 'Value', 'Metric', 'Value'],
        ['Departments / segments', str(len(vlsm)),
         'VLANs defined', str(len(vlan) or len(vlsm))],
        ['Total usable hosts', str(total_usable),
         'IPv4 assignments', str(len(ipv4))],
        ['Routers', str(len(routers)),
         'Switches', str(len(switches))],
        ['Internet / WAN link', 'Enabled' if internet_on else 'Disabled',
         'Routing protocol', str(project.get('routing_protocol') or 'OSPF')],
        ['Base network', str(project.get('base_network') or 'N/A'),
         'Project status', str(project.get('status') or 'Draft')],
        ['Validation passed', str(checks_pass),
         'Validation failed', str(checks_fail)],
    ]
    mt1 = Table(metrics1, colWidths=[1.7*inch, 1.4*inch, 1.7*inch, 1.4*inch])
    mt1.setStyle(_table_style())
    story.append(mt1)
    story.append(Paragraph('Table 1 – Key design metrics', S['RCaption']))

    # Row 2 – addressing highlights
    metrics2 = [
        ['Addressing highlight', 'Detail'],
        ['Largest segment',
         f'{largest.get("dept", "—")}  ·  {largest.get("network", "")}{largest.get("prefix", "")}  ·  '
         f'{largest.get("usable", "—")} hosts  ·  VLAN {largest.get("vlan_id", "—")}'],
        ['Smallest segment',
         f'{smallest.get("dept", "—")}  ·  {smallest.get("network", "")}{smallest.get("prefix", "")}  ·  '
         f'{smallest.get("usable", "—")} hosts  ·  VLAN {smallest.get("vlan_id", "—")}'],
        ['VLAN ID range',
         f'{min(vlan_ids)} – {max(vlan_ids)}' if vlan_ids else '—'],
        ['Device inventory',
         'Routers: ' + ', '.join(r.get('name', 'R') for r in routers)
         + '  |  Switches: ' + ', '.join(s.get('name', 'S') for s in switches)],
    ]
    mt2 = Table(metrics2, colWidths=[1.7*inch, 4.5*inch])
    mt2.setStyle(_table_style())
    story.append(mt2)
    story.append(Paragraph('Table 2 – Addressing and inventory highlights', S['RCaption']))

    # 2. Project Summary
    story.append(Paragraph('2. Project Summary', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    summary_data = [
        ['Field', 'Value'],
        ['Company', project.get('company_name', '')],
        ['Project', project.get('project_name', '')],
        ['Base Network', project.get('base_network', '')],
        ['Routing Protocol', project.get('routing_protocol', '')],
        ['Primary Router', project.get('router_name', 'R1')],
        ['Primary Switch', project.get('switch_name', 'S1')],
        ['Routers (count)', str(len(routers))],
        ['Switches (count)', str(len(switches))],
        ['Internet / WAN', 'Enabled' if internet_on else 'Disabled'],
        ['Departments', str(len(vlsm))],
        ['Total usable hosts', str(total_usable)],
        ['VLANs', str(len(vlan) or len(vlsm))],
        ['IPv4 assignments', str(len(ipv4))],
        ['IPv6 plan entries', str(len(generated.get('ipv6') or []))],
        ['Validation (pass / fail)', f'{checks_pass} / {checks_fail}'],
        ['Status', project.get('status', '')],
    ]
    t = Table(summary_data, colWidths=[2.2 * inch, 4.0 * inch])
    t.setStyle(_table_style())
    story.append(t)
    if generated.get('network_summary'):
        story.append(Spacer(1, 8))
        story.append(Paragraph(generated['network_summary'], S['RBody']))

    # 3. Topology
    story.append(Paragraph('3. Network Topology Diagram', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    story.append(Paragraph(
        'Hierarchical view of the planned network: Internet edge, routers, switches, and department segments.',
        S['RBody'],
    ))
    try:
        story.append(DiagramFlowable(
            _draw_topology_diagram(project, generated),
            'Figure 3 – Hierarchical network topology',
        ))
    except Exception as e:
        story.append(Paragraph(f'[Topology diagram unavailable: {e}]', S['RBody']))

    # 3. UML
    story.append(PageBreak())
    story.append(Paragraph('4. UML Diagrams', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    story.append(Paragraph(
        'Use-case, activity, and component views describing the network design and NETRIX workflow.',
        S['RBody'],
    ))
    try:
        story.append(DiagramFlowable(_draw_uml_usecase(project), 'Figure 4.1 – Use-case diagram'))
        story.append(Spacer(1, 8))
        story.append(DiagramFlowable(_draw_uml_activity(project), 'Figure 4.2 – Activity flow'))
        story.append(Spacer(1, 8))
        story.append(DiagramFlowable(
            _draw_uml_class(project, generated),
            'Figure 4.3 – Component / class view',
        ))
    except Exception as e:
        story.append(Paragraph(f'[UML diagrams unavailable: {e}]', S['RBody']))

    # 4. VLSM
    story.append(PageBreak())
    story.append(Paragraph('5. VLSM Subnet Table', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    vlsm = generated.get('vlsm', [])
    if vlsm:
        rows = [['Department', 'Network', 'Prefix', 'Host Range', 'Usable', 'VLAN']]
        for v in vlsm:
            rows.append([
                str(v.get('dept', '')),
                str(v.get('network', '')),
                str(v.get('prefix', '')),
                str(v.get('range', '')),
                str(v.get('usable', '')),
                str(v.get('vlan_id', '')),
            ])
        t = Table(rows, colWidths=[1.15*inch, 1.15*inch, 0.6*inch, 1.7*inch, 0.6*inch, 0.55*inch])
        t.setStyle(_table_style())
        story.append(t)
    else:
        story.append(Paragraph('No VLSM data available.', S['RBody']))

    # 5. VLAN
    story.append(Paragraph('6. VLAN Allocation', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    vlan = generated.get('vlan', [])
    if vlan:
        rows = [['VLAN ID', 'Department', 'Network / Prefix']]
        for v in vlan:
            rows.append([str(v.get('id', '')), str(v.get('dept', '')), str(v.get('network', ''))])
        t = Table(rows, colWidths=[1.0 * inch, 2.2 * inch, 3.0 * inch])
        t.setStyle(_table_style())
        story.append(t)

    # 6. IPv4
    story.append(Paragraph('7. IPv4 Address Plan', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    ipv4 = generated.get('ipv4', [])
    if ipv4:
        rows = [['Device', 'Interface', 'IP Address', 'Mask']]
        for row in ipv4[:45]:
            rows.append([
                str(row.get('device', '')),
                str(row.get('interface', '')),
                str(row.get('ip', '')),
                str(row.get('mask', '')),
            ])
        t = Table(rows, colWidths=[1.5 * inch, 1.8 * inch, 1.5 * inch, 1.4 * inch])
        t.setStyle(_table_style())
        story.append(t)

    # 7. Validation
    
    # 6b. IPv6 Address Plan (always rendered when data or dual/ipv6 design present)
    story.append(Paragraph('8. IPv6 Address Plan (ULA)', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    ipv6 = list(generated.get('ipv6') or [])
    design = generated.get('design') or {}
    ip_ver = (design.get('ip_version') or 'ipv4').lower()

    # Fallback: parse IPv6 from router config if JSON empty
    if not ipv6 and generated.get('router_config'):
        import re as _re
        cfg = generated.get('router_config') or ''
        for m in _re.finditer(
            r'interface\s+(\S+)[\s\S]*?ipv6 address\s+([0-9a-fA-F:]+)/(\d+)',
            cfg,
        ):
            iface, addr, plen = m.group(1), m.group(2), m.group(3)
            # network = address with interface id zeroed roughly
            net = addr.rsplit(':', 1)[0] + '::'
            ipv6.append({
                'dept': iface,
                'vlan_id': '',
                'network': f'{net}/{plen}',
                'gateway': f'{addr}',
                'range': f'{addr} – (subnet /{plen})',
                'prefix': f'/{plen}',
            })

    story.append(Paragraph(
        f'<b>Design IP version:</b> {ip_ver.upper()} &nbsp;|&nbsp; '
        f'<b>IPv6 rows:</b> {len(ipv6)} &nbsp;|&nbsp; '
        f'<b>ULA base:</b> <font face="Courier">fd00:9e71::/32</font>',
        S['RBody'],
    ))
    story.append(Spacer(1, 4))

    if ipv6:
        story.append(Paragraph(
            'Unique Local Addresses (ULA, RFC 4193 style) are allocated as <font face="Courier">/64</font> '
            'subnets per department VLAN and optional WAN link. Gateway addresses use host interface '
            'identifier <font face="Courier">::1</font>. ULA space is not advertised on the public Internet.',
            S['RBody'],
        ))
        story.append(Spacer(1, 6))
        rows6 = [['Department / Link', 'VLAN', 'Network prefix', 'Gateway', 'Host range', 'Prefix']]
        for row in ipv6[:50]:
            rows6.append([
                str(row.get('dept', '')),
                str(row.get('vlan_id') if row.get('vlan_id') not in (None, '', 0) else '—'),
                str(row.get('network', '')),
                str(row.get('gateway', '')),
                str(row.get('range', '')),
                str(row.get('prefix', '/64')),
            ])
        t6 = Table(rows6, colWidths=[1.25*inch, 0.55*inch, 1.55*inch, 1.25*inch, 1.35*inch, 0.55*inch])
        t6.setStyle(_table_style())
        story.append(t6)
        story.append(Paragraph('Table 7 – IPv6 ULA allocation', S['RCaption']))
        story.append(Spacer(1, 6))
        # Dual-stack mapping when both present
        vlsm_rows = generated.get('vlsm') or []
        if ip_ver == 'dual' and vlsm_rows:
            story.append(Paragraph('<b>Dual-stack mapping (IPv4 ↔ IPv6 by department)</b>', S['RH2']))
            map_rows = [['Department', 'VLAN', 'IPv4 network', 'IPv6 network', 'IPv6 gateway']]
            v6_by_dept = {str(r.get('dept', '')).lower(): r for r in ipv6}
            for v in vlsm_rows[:30]:
                dept = v.get('dept', '')
                v6 = v6_by_dept.get(str(dept).lower()) or {}
                map_rows.append([
                    str(dept),
                    str(v.get('vlan_id', '')),
                    str(v.get('network', '')),
                    str(v6.get('network', '—')),
                    str(v6.get('gateway', '—')),
                ])
            tm = Table(map_rows, colWidths=[1.3*inch, 0.55*inch, 1.5*inch, 1.6*inch, 1.5*inch])
            tm.setStyle(_table_style())
            story.append(tm)
            story.append(Paragraph('Table 7b – Dual-stack department mapping', S['RCaption']))
            story.append(Spacer(1, 6))
        story.append(Paragraph(
            '<b>Implementation notes.</b> Enable <font face="Courier">ipv6 unicast-routing</font> on routers. '
            'Apply <font face="Courier">ipv6 address &lt;gateway&gt;/64</font> on each SVI or subinterface. '
            'For dual-stack, keep VLAN IDs identical for IPv4 and IPv6 segments. Prefer OSPFv3 or static '
            'IPv6 routes for inter-VLAN forwarding in lab environments.',
            S['RBody'],
        ))
    else:
        story.append(Paragraph(
            f'No IPv6 plan rows are stored for this project (recorded IP version: <b>{ip_ver}</b>). '
            f'To include a full IPv6 ULA table, open <b>New Project</b>, set IP Version to '
            f'<b>IPv6 only</b> or <b>Dual-stack</b>, then regenerate and download the PDF again.',
            S['RBody'],
        ))


    # 7. Validation  (renumber note in heading if needed)

    story.append(Paragraph('9. Validation Results', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    for v in generated.get('validation', []):
        mark = 'PASS' if v.get('pass') else 'FAIL'
        story.append(Paragraph(f'<b>[{mark}]</b> {v.get("msg", "")}', S['RBody']))

    # 8–9 Configs
    story.append(Paragraph('10. Router Configuration (excerpt)', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    cfg = generated.get('router_config', '') or ''
    story.append(Preformatted(
        cfg[:1800] + ('\n...' if len(cfg) > 1800 else ''),
        S['RCode'],
    ))

    story.append(Paragraph('11. Switch Configuration (excerpt)', S['RH1']))
    story.append(HRFlowable(width='100%', thickness=0.8, color=GREY_LINE, spaceAfter=8))
    cfg = generated.get('switch_config', '') or ''
    story.append(Preformatted(
        cfg[:1500] + ('\n...' if len(cfg) > 1500 else ''),
        S['RCode'],
    ))

    # Closing note
    story.append(Spacer(1, 16))
    story.append(HRFlowable(width='100%', thickness=1, color=NAVY, spaceAfter=6))
    story.append(Paragraph(
        'This document was generated by the NETRIX Enterprise Network Planning Framework. '
        'Validate the design in Cisco Packet Tracer using the exported lab guide and device configurations.',
        S['RCaption'],
    ))

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    buffer.seek(0)
    return buffer.read()


# ─── Excel / Word / CSV (unchanged structure, minor polish) ─────────────────

def generate_excel_report(project: dict, generated: dict) -> bytes:
    wb = Workbook()
    header_fill = PatternFill('solid', fgColor='0F2744')
    header_font = Font(bold=True, color='FFFFFF', name='Tahoma', size=11)
    body_font = Font(name='Tahoma', size=11, color='000000')

    ws = wb.active
    ws.title = 'Summary'
    ws.append(['Field', 'Value'])
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
    for label, val in [
        ('Company', project.get('company_name')),
        ('Project', project.get('project_name')),
        ('Base Network', project.get('base_network')),
        ('Routing', project.get('routing_protocol')),
        ('Status', project.get('status')),
        ('Summary', generated.get('network_summary', '')),
    ]:
        ws.append([label, val])
        for cell in ws[ws.max_row]:
            cell.font = body_font
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 55

    ws2 = wb.create_sheet('VLSM')
    headers = ['Department', 'Network', 'Prefix', 'Mask', 'Host Range', 'Usable Hosts', 'VLAN ID']
    ws2.append(headers)
    for cell in ws2[1]:
        cell.fill = header_fill
        cell.font = header_font
    for v in generated.get('vlsm', []):
        ws2.append([v.get('dept'), v.get('network'), v.get('prefix'), v.get('mask', ''),
                    v.get('range'), v.get('usable'), v.get('vlan_id', '')])
        for cell in ws2[ws2.max_row]:
            cell.font = body_font
    for col in ws2.columns:
        ws2.column_dimensions[col[0].column_letter].width = 18

    ws3 = wb.create_sheet('IPv4 Plan')
    ws3.append(['Device', 'Interface', 'IP Address', 'Subnet Mask'])
    for cell in ws3[1]:
        cell.fill = header_fill
        cell.font = header_font
    for d in generated.get('ipv4', []):
        ws3.append([d.get('device'), d.get('interface'), d.get('ip'), d.get('mask', '')])
        for cell in ws3[ws3.max_row]:
            cell.font = body_font
    for col in ws3.columns:
        ws3.column_dimensions[col[0].column_letter].width = 20

    ws4 = wb.create_sheet('VLANs')
    ws4.append(['VLAN ID', 'Department', 'Network'])
    for cell in ws4[1]:
        cell.fill = header_fill
        cell.font = header_font
    for v in generated.get('vlan', []):
        ws4.append([v.get('id'), v.get('dept'), v.get('network')])
        for cell in ws4[ws4.max_row]:
            cell.font = body_font

    ws5 = wb.create_sheet('Router Config')
    for line in (generated.get('router_config') or '').splitlines():
        ws5.append([line])
    ws5.column_dimensions['A'].width = 80

    ws6 = wb.create_sheet('Switch Config')
    for line in (generated.get('switch_config') or '').splitlines():
        ws6.append([line])
    ws6.column_dimensions['A'].width = 80

    ws7 = wb.create_sheet('IPv6 Plan')
    ws7.append(['Department / Link', 'VLAN', 'Network prefix', 'Gateway', 'Host range', 'Prefix length'])
    for cell in ws7[1]:
        cell.fill = header_fill
        cell.font = header_font
    for row in (generated.get('ipv6') or []):
        ws7.append([
            row.get('dept', ''),
            row.get('vlan_id', ''),
            row.get('network', ''),
            row.get('gateway', ''),
            row.get('range', ''),
            row.get('prefix', ''),
        ])
        for cell in ws7[ws7.max_row]:
            cell.font = body_font
    for col in ws7.columns:
        ws7.column_dimensions[col[0].column_letter].width = 22

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_word_report(project: dict, generated: dict) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Tahoma'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)

    title = doc.add_heading('NETRIX – Enterprise Network Planning Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Tahoma'
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')
    for run in p.runs:
        run.font.name = 'Tahoma'
        run.font.size = Pt(12)

    doc.add_heading('1. Executive Summary', level=1)
    for para in _build_executive_summary(project, generated):
        plain = para.replace('<b>', '').replace('</b>', '')
        doc.add_paragraph(plain)

    # Key metrics in Word
    vlsm = generated.get('vlsm') or []
    vlan = generated.get('vlan') or []
    ipv4 = generated.get('ipv4') or []
    validation = generated.get('validation') or []
    routers = project.get('routers') or [{'name': project.get('router_name') or 'R1'}]
    switches = project.get('switches') or [{'name': project.get('switch_name') or 'S1'}]
    total_usable = sum(int(v.get('usable') or 0) for v in vlsm)
    checks_pass = sum(1 for v in validation if v.get('pass'))
    checks_fail = len(validation) - checks_pass
    doc.add_heading('Key Metrics', level=2)
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Table Grid'
    metrics_table.rows[0].cells[0].text = 'Metric'
    metrics_table.rows[0].cells[1].text = 'Value'
    for label, val in [
        ('Departments / segments', str(len(vlsm))),
        ('Total usable hosts', str(total_usable)),
        ('VLANs', str(len(vlan) or len(vlsm))),
        ('IPv4 assignments', str(len(ipv4))),
        ('Routers', str(len(routers))),
        ('Switches', str(len(switches))),
        ('Routing protocol', str(project.get('routing_protocol') or 'OSPF')),
        ('Base network', str(project.get('base_network') or 'N/A')),
        ('Internet / WAN', 'Enabled' if (project.get('internet') or {}).get('enabled', True) else 'Disabled'),
        ('Validation passed', str(checks_pass)),
        ('Validation failed', str(checks_fail)),
        ('Project status', str(project.get('status') or 'Draft')),
    ]:
        row = metrics_table.add_row().cells
        row[0].text = label
        row[1].text = val

    doc.add_heading('2. Project Summary', level=1)
    for label, key in [
        ('Company', 'company_name'), ('Project', 'project_name'),
        ('Base Network', 'base_network'), ('Routing', 'routing_protocol'),
        ('Router', 'router_name'), ('Switch', 'switch_name'),
    ]:
        doc.add_paragraph(f'{label}: {project.get(key, "")}')
    if generated.get('network_summary'):
        doc.add_paragraph(generated['network_summary'])

    doc.add_heading('3. Network Topology (description)', level=1)
    depts = [v.get('dept') for v in generated.get('vlsm', [])]
    doc.add_paragraph(
        f'Topology: Internet → {project.get("router_name", "R1")} → '
        f'{project.get("switch_name", "S1")} → Departments ({", ".join(depts) or "N/A"}).'
    )

    doc.add_heading('4. VLSM Subnet Table', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, h in enumerate(['Department', 'Network', 'Prefix', 'Range', 'Usable']):
        table.rows[0].cells[i].text = h
    for v in generated.get('vlsm', []):
        row = table.add_row().cells
        row[0].text = str(v.get('dept', ''))
        row[1].text = str(v.get('network', ''))
        row[2].text = str(v.get('prefix', ''))
        row[3].text = str(v.get('range', ''))
        row[4].text = str(v.get('usable', ''))

    doc.add_heading('5. VLAN Allocation', level=1)
    for v in generated.get('vlan', []):
        doc.add_paragraph(f"VLAN {v.get('id')}: {v.get('dept')} → {v.get('network')}", style='List Bullet')

    doc.add_heading('6. IPv6 Address Plan (ULA)', level=1)
    ipv6 = generated.get('ipv6') or []
    design = generated.get('design') or {}
    doc.add_paragraph(f"IP version: {(design.get('ip_version') or 'ipv4')}")
    if ipv6:
        table6 = doc.add_table(rows=1, cols=5)
        hdr = table6.rows[0].cells
        for i, h in enumerate(['Department', 'VLAN', 'Network', 'Gateway', 'Range']):
            hdr[i].text = h
        for row in ipv6[:40]:
            cells = table6.add_row().cells
            cells[0].text = str(row.get('dept', ''))
            cells[1].text = str(row.get('vlan_id', '') or '—')
            cells[2].text = str(row.get('network', ''))
            cells[3].text = str(row.get('gateway', ''))
            cells[4].text = str(row.get('range', ''))
    else:
        doc.add_paragraph('No IPv6 plan rows. Select IPv6 or Dual-stack and regenerate.')

    doc.add_heading('7. Router Configuration', level=1)
    p = doc.add_paragraph()
    run = p.add_run((generated.get('router_config') or '')[:2000])
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('7. Switch Configuration', level=1)
    p = doc.add_paragraph()
    run = p.add_run((generated.get('switch_config') or '')[:1500])
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_csv_vlsm(vlsm_list) -> str:
    lines = ['Department,Network,Prefix,Mask,Host Range,Usable Hosts,VLAN ID']
    for v in vlsm_list:
        lines.append(','.join([
            str(v.get('dept', '')), str(v.get('network', '')), str(v.get('prefix', '')),
            str(v.get('mask', '')), f'"{v.get("range", "")}"',
            str(v.get('usable', '')), str(v.get('vlan_id', '')),
        ]))
    return '\n'.join(lines)


def generate_csv_ipv4(ipv4_list) -> str:
    lines = ['Device,Interface,IP Address,Subnet Mask']
    for d in ipv4_list:
        lines.append(','.join([
            str(d.get('device', '')), str(d.get('interface', '')),
            str(d.get('ip', '')), str(d.get('mask', '')),
        ]))
    return '\n'.join(lines)
